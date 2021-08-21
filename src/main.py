import os
import random
import secrets
import string

import docx
import json
import wave

from vosk import Model, KaldiRecognizer
import vosk

import moviepy.editor as mp
from pydub import AudioSegment


class SoundMixin:
    """
        class for return text from audio file
        example to use:
            result = SoundMixin('/{}.wav'.format(audio)).get_data()
    """

    @staticmethod
    def _sliced_sound(audio_file):
        song = AudioSegment.from_wav(audio_file)
        return song[0:120000]

    @staticmethod
    def get_data(model: vosk.Model, audio: str) -> str:
        """
            Recognize Russian voice in wav
        """

        wave_audio_file = wave.open(audio, "rb")
        offline_recognizer = KaldiRecognizer(model, 70000)
        data = wave_audio_file.readframes(wave_audio_file.getnframes())

        offline_recognizer.AcceptWaveform(data)
        recognized_data = json.loads(offline_recognizer.Result())["text"]
        return recognized_data


class VideoConvertor:
    """
        class for converting format from video to audio
        example to use:
            audio = VideoConvertor('test.mp4').to_convert(0, 21)
    """

    def __init__(self):
        self.BASE_DIR: str = os.path.dirname(os.path.abspath(__file__))

    @staticmethod
    def _take_len(video, start: int):
        return video.subclip(start)

    def to_convert(self, video, start: int) -> str:
        video = mp.VideoFileClip(video)
        audio = str(random.randint(1, 1000000000))
        self._take_len(video, start).audio.write_audiofile(
            "{}.wav".format(audio)
        )
        return audio


class MainManager(VideoConvertor, SoundMixin):

    @staticmethod
    def _set_name():
        letters_and_digits = string.ascii_letters + string.digits
        return ''.join(secrets.choice(letters_and_digits) for i in range(32))

    def get_text_from_video(self, video):
        audio = self.to_convert(video, 0)
        return self.get_data(
            model=Model(self.BASE_DIR + "/vosk-model-small-ru-0.15"),
            audio=self.BASE_DIR + '/{}.wav'.format(f'{audio}')
        )

    @staticmethod
    def create_doc(text):
        document = docx.Document()
        par1 = document.add_paragraph()
        par1.add_run(text)
        #   path_file = self.BASE_DIR + '/{}.docx'.format(self._set_name())
        #   document.save(path_file)

        return document
