import os
import random
import secrets
import string

import speech_recognition as speech_recog
from speech_recognition import Recognizer as recog

import docx

import moviepy.editor as mp


class SoundMixin:
    """
        class for return text from audio file
        example to use:
            result = SoundMixin('/{}.wav'.format(audio)).get_data()
    """

    @staticmethod
    def _get_content(audio_file):
        with audio_file as audio_file:
            audio_content = recog().record(audio_file)

        return audio_content

    def get_data(self, audio_file) -> str:
        return recog().recognize_google(self._get_content(audio_file), language='ru-RU')


class VideoConvertor:
    """
        class for converting format from video to audio
        example to use:
            audio = VideoConvertor('test.mp4').to_convert(0, 21)
    """

    def __init__(self):
        self.BASE_DIR: str = os.path.dirname(os.path.abspath(__file__))

    @staticmethod
    def _take_len(video, start: int):
        return video.subclip(start)

    def to_convert(self, video, start: int) -> str:
        video = mp.VideoFileClip(video)
        audio = str(random.randint(1, 1000000000))
        self._take_len(video, start).audio.write_audiofile(
            "{}.wav".format(audio)
        )
        return audio


class MainManager(VideoConvertor, SoundMixin):

    @staticmethod
    def _set_name():
        letters_and_digits = string.ascii_letters + string.digits
        return ''.join(secrets.choice(letters_and_digits) for i in range(32))

    def get_text_from_video(self, video):
        audio = self.to_convert(video, 0)
        return self.get_data(speech_recog.AudioFile(self.BASE_DIR + '/{}.wav'.format(audio)))

    @staticmethod
    def create_doc(text):
        document = docx.Document()
        par1 = document.add_paragraph()
        par1.add_run(text)
        #   path_file = self.BASE_DIR + '/{}.docx'.format(self._set_name())
        #   document.save(path_file)

        return document
